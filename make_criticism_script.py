import sys
import subprocess

try:
    import docx
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    import docx
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = docx.Document()

def h(text, level=2): return doc.add_heading(text, level=level)
def b(text): return doc.add_paragraph(text)
def bullet(bold_label, text):
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(bold_label + ": ").bold = True
    p.add_run(text)
def divider(): doc.add_paragraph("═" * 65)
def tag(text):
    p = doc.add_paragraph()
    p.add_run(f"▸ {text}").bold = True
    return p

# ═══════════════════════════════════════════════════
#  TITLE
# ═══════════════════════════════════════════════════
t = doc.add_heading("ECOZONE Fingerprint Attendance System", 0)
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
s = doc.add_heading("Panelist Criticism Rebuttal Script  |  3 Members", level=1)
s.alignment = WD_ALIGN_PARAGRAPH.CENTER
info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
info.add_run("Miko R. Zamora   |   Razel T. Herodias   |   Kevin Trazie C. Librero\n")
info.add_run("ZAMBOECOZONE – MIS Division  |  BSCS-2A  |  WMSU\n").bold = True

divider()
note = doc.add_paragraph()
note.add_run("PURPOSE: ").bold = True
note.add_run(
    "This script prepares all three members to handle the toughest, "
    "most hostile criticism a panelist can throw. Every attack is met with a confident, "
    "detailed, technically accurate rebuttal. Assigned member is marked per criticism."
)
divider()
doc.add_page_break()

# ═══════════════════════════════════════════════════
#  SECTION 1 — SECURITY CRITICISMS
# ═══════════════════════════════════════════════════
h("SECTION 1: Security Criticisms", level=1)
doc.add_paragraph("Owner: KEVIN — handles all security-related attacks")
divider()

# --- Criticism 1 ---
tag("CRITICISM 1 — \"You're storing fingerprints as plain hex text. That's not secure at all.\"")
b("SEVERITY: This is the #1 attack panelists will use. Be ready.")
b("REBUTTAL (Kevin):")
b(
    "\"With respect, that is a misunderstanding of both our data and the threat model. "
    "What we store is not a fingerprint — it is an ANSI 378 Minutiae Template. It is a mathematical "
    "abstraction of the structural geometry of ridge patterns, encoded in binary and then expressed as a hex string. "
    "There is no known algorithm that can reverse-engineer an ANSI 378 template back into a fingerprint image — "
    "the conversion is lossy and one-directional by design. "
    "Compare this to traditional password hashing: we hash passwords to prevent dictionary attacks against "
    "recognizable data. A Minutiae template has no human-recognizable pattern even in its raw form. "
    "Furthermore, our system never stores the raw pixel image of the fingerprint — only the extracted template. "
    "This is exactly the approach used in government biometric ID systems compliant with ISO/IEC 19794-2. "
    "So the data is both non-reversible and non-recognizable — it is secure by design.\""
)

# --- Criticism 2 ---
tag("CRITICISM 2 — \"You should be using bcrypt or AES encryption on the biometric data.\"")
b("REBUTTAL (Kevin):")
b(
    "\"Bcrypt is a password hashing algorithm — it is specifically designed for fixed, short, deterministic inputs "
    "like text passwords. A biometric template is none of those things. It is a variable-length binary structure. "
    "Applying bcrypt would produce a completely different output even for a slightly rotated legitimate scan of the "
    "same finger, making verification impossible. "
    "AES encryption is a valid consideration, but it would require decrypting every stored template before comparison, "
    "which (1) reintroduces the raw data in memory during matching, creating a different attack surface, and "
    "(2) is architecturally unnecessary because the stored templates are already non-interpretable without the SDK. "
    "The ANSI 378 standard is the established, internationally recognized biometric security approach — "
    "used by the FBI, Interpol, and national ID systems worldwide.\""
)

# --- Criticism 3 ---
tag("CRITICISM 3 — \"Your system is vulnerable to SQL injection.\"")
b("REBUTTAL (Kevin):")
b(
    "\"Every single database query in our system uses MySQLi prepared statements with bound parameters. "
    "There is no instance of a raw string being concatenated into a SQL query. "
    "The nickname input is bound as a string parameter. The fingerprint hex data never touches SQL directly — "
    "it is written to a temp file and passed to the C++ executable as a command-line argument, not injected into SQL. "
    "If you would like, I can show you the exact prepared statement in enroll.php line 188 right now.\""
)

# --- Criticism 4 ---
tag("CRITICISM 4 — \"What stops someone from intercepting the hex data and replaying it?\"")
b("REBUTTAL (Kevin):")
b(
    "\"A replay attack — submitting a captured hex string directly — is addressed in two ways. "
    "First, during enrollment, the CHECK mode of scanner.exe compares the submitted hex against all existing "
    "database records. If that hex was already registered, it is blocked as a duplicate. "
    "Second, and more importantly: a hex string that passes duplicate detection during enrollment still cannot "
    "be used for attendance fraud, because the verification mode requires a LIVE SCAN from the physical device. "
    "The verify mode calls scanner.exe verify — which opens the hardware scanner and requires a physical finger "
    "to be placed on the sensor. There is no API endpoint that accepts a raw hex value for attendance logging — "
    "only real-time hardware capture triggers verification.\""
)

doc.add_page_break()

# ═══════════════════════════════════════════════════
#  SECTION 2 — TECHNICAL / ARCHITECTURE CRITICISMS
# ═══════════════════════════════════════════════════
h("SECTION 2: Technical & Architecture Criticisms", level=1)
doc.add_paragraph("Owner: MIKO — handles all architecture and design-related attacks")
divider()

# --- Criticism 5 ---
tag("CRITICISM 5 — \"Why did you use shell_exec() in PHP? That's a serious security risk.\"")
b("REBUTTAL (Miko):")
b(
    "\"Shell_exec() can be a risk when used carelessly — specifically when user-supplied input is directly "
    "concatenated into the shell command. We are aware of this. In our implementation, the shell command is "
    "constructed entirely by server-side PHP logic. The only user-supplied value that appears in the command "
    "is the fingerprint hex string in CHECK mode — and that value was itself already generated by our own scanner.exe "
    "in a prior server-side call, not typed by a user. "
    "We also validated that the scanner executable path is resolved using PHP's realpath(), which prevents "
    "path traversal attacks. Safe-to-execute arguments are limited and server-controlled. "
    "In a production deployment, this would be additionally hardened by running the PHP process under a restricted "
    "service account with no shell access to sensitive system resources.\""
)

# --- Criticism 6 ---
tag("CRITICISM 6 — \"Your system only works locally on one machine. It has no scalability.\"")
b("REBUTTAL (Miko):")
b(
    "\"You are correct that in its current form, the system is designed for deployment on a single Windows "
    "workstation with a connected scanner — which is appropriate for ZAMBOECOZONE's MIS Division office environment "
    "where a single attendance station is the expected deployment context. "
    "However, the architecture is intentionally modular. The C++ scanner executable is decoupled from the web backend "
    "by design. For multi-station scaling, each attendance terminal would run its own scanner.exe, and all terminals "
    "would point to a single centralized MySQL database server. The PHP API is stateless and does not hold "
    "in-memory session data — it was specifically designed this way to support multiple simultaneous client connections "
    "in a future network deployment.\""
)

# --- Criticism 7 ---
tag("CRITICISM 7 — \"Why not use a commercial off-the-shelf HRIS instead of custom-building this?\"")
b("REBUTTAL (Miko):")
b(
    "\"ZAMBOECOZONE, as a government agency, operates under specific IT procurement and data privacy constraints. "
    "Commercial biometric HRIS platforms often store biometric data on vendor cloud servers — which would violate "
    "RA 10173 requirements for data localization and consent for a government entity of this nature. "
    "Our system keeps all biometric data entirely on-premises, on servers under MIS Division control. "
    "Additionally, the objective of an OJT project is precisely to develop hands-on engineering competency. "
    "Building a custom solution gave our team direct exposure to hardware SDK integration, "
    "secure API design, and database architecture — which no off-the-shelf product could provide.\""
)

# --- Criticism 8 ---
tag("CRITICISM 8 — \"The C++ executable is a black box to the PHP side. How do you handle errors?\"")
b("REBUTTAL (Miko):")
b(
    "\"Our C++ executable follows a strict output protocol. Every output begins with a known prefix: "
    "ERROR: for any failure (with a human-readable message), MATCH:[id] for a successful verification, "
    "NOMATCH for no match found, or a raw hex string for an enrollment. "
    "On the PHP side, every shell_exec() return value is checked using strpos() for the ERROR: prefix before "
    "any further processing. If an error is returned, PHP immediately returns a JSON error response to the frontend "
    "with the scanner's error message. No partial data is ever committed to the database on an error. "
    "This protocol forms a clear, reliable contract between the two layers.\""
)

doc.add_page_break()

# ═══════════════════════════════════════════════════
#  SECTION 3 — SCOPE / OJT CRITICISMS
# ═══════════════════════════════════════════════════
h("SECTION 3: Scope & OJT-Level Criticisms", level=1)
doc.add_paragraph("Owner: RAZEL — handles all scope, project relevance, and OJT-related attacks")
divider()

# --- Criticism 9 ---
tag("CRITICISM 9 — \"This is too complex for a 2nd-year BSCS OJT project. Did you really build all of this?\"")
b("REBUTTAL (Razel):")
b(
    "\"Yes, every line of code in this system was written by our team, and our portfolio documents the week-by-week "
    "progression that proves it. We did not start with a finished product — we started from scratch in Week 3, "
    "reached 40-50% completion by Week 5 when our mentor evaluated and gave feedback, and iteratively improved "
    "through Weeks 9 and 10. "
    "The source code is on our GitHub repository with a commit history showing dated contributions from all three members. "
    "Complexity is not an argument against authenticity — it is an argument for how seriously we took this OJT. "
    "We also had the guidance of our MIS Division supervisors to ensure correctness.\""
)

# --- Criticism 10 ---
tag("CRITICISM 10 — \"What happens if the scanner breaks? Your entire attendance system goes down.\"")
b("REBUTTAL (Razel):")
b(
    "\"That is a valid operational concern. Our system is designed so that the scanner failure is graceful — "
    "if scanner.exe cannot detect the device, it returns ERROR:No scanner found, which the PHP API relays "
    "to the frontend as a clear error message. The web interface and database remain fully operational. "
    "Attendance records already stored are always accessible. "
    "For production resilience — which would be a Phase 2 recommendation to MIS — we would propose keeping a "
    "backup scanner unit and implementing a manual override log for hardware downtime events. "
    "This is consistent with how any hardware-dependent enterprise system handles peripheral failure.\""
)

# --- Criticism 11 ---
tag("CRITICISM 11 — \"You didn't implement a login system. Anyone can open the enrollment page.\"")
b("REBUTTAL (Razel):")
b(
    "\"That is correct, and it is a deliberate scope decision for this OJT phase. ZAMBOECOZONE's MIS Division "
    "indicated that the enrollment workstation would be operated only by authorized IT staff in a controlled office environment. "
    "The priority for this phase was proving the core biometric pipeline — hardware integration, duplicate detection, "
    "and data storage — not building a full access control layer around it. "
    "That said, we fully acknowledge that a production deployment would require admin authentication — "
    "specifically session-based PHP login with role-based access control — before any enrollment or admin action. "
    "This is documented as a future enhancement in our system documentation.\""
)

# --- Criticism 12 ---
tag("CRITICISM 12 — \"What is the actual business value of this system to ZAMBOECOZONE?\"")
b("REBUTTAL (Razel):")
b(
    "\"ZAMBOECOZONE is a government economic zone authority that employs hundreds of staff across multiple divisions. "
    "Manual attendance tracking is both time-consuming and legally auditable — errors or buddy-punching can create "
    "payroll disputes and HR compliance issues. "
    "Our system eliminates buddy-punching by design — it is physically impossible to clock in for another employee "
    "because the system requires the registered finger on the registered hardware. "
    "It also creates a verifiable, timestamped digital attendance record that can be queried, exported, and audited "
    "at any time — unlike handwritten logbooks. "
    "For MIS Division's operational context, this is a direct improvement in data integrity, staff accountability, "
    "and administrative efficiency.\""
)

doc.add_page_break()

# ═══════════════════════════════════════════════════
#  MASTER CONFIDENCE GUIDE
# ═══════════════════════════════════════════════════
h("MASTER CONFIDENCE RULES — READ BEFORE YOUR DEFENSE", level=1)
divider()

rules = [
    ("Never say \"I don't know\"",
     "Say: \"That is a great question. That specific scenario was outside our OJT scope, but our architectural approach would handle it by [explain]. We'd be happy to discuss that as a future enhancement.\""),
    ("Never apologize for limitations",
     "Reframe every limitation as a deliberate scope decision: \"That was a conscious design boundary for this OJT phase, and here is the rationale...\""),
    ("If a panelist interrupts you",
     "Say: \"Understood — let me address that precisely.\" Then give a short, direct answer before continuing."),
    ("If you are asked something the other member owns",
     "Say: \"[Name] is our technical lead on that — [Name], would you like to add more detail?\" Then hand off smoothly."),
    ("If a panelist says \"that's insecure\"",
     "Ask them to clarify the specific attack vector: \"Could you specify the threat model you have in mind? We have addressed several vectors including replay attacks, SQL injection, and data reversal — I want to be sure I answer the exact concern.\""),
    ("If you are wrong",
     "Never argue. Say: \"Thank you for that correction. You are right that in a production environment, [acknowledge the point]. During our OJT, we prioritized [x] first, but that is a valid enhancement to document.\""),
]

for title_text, content in rules:
    b_title = doc.add_paragraph()
    b_title.add_run(f"▸ RULE: {title_text}").bold = True
    doc.add_paragraph(f"   → {content}")

output_path = r'C:\Users\Administrator\Desktop\ECOZONE_Criticism_Rebuttal.docx'
doc.save(output_path)
print(f"SUCCESS: DOCX saved to {output_path}")
