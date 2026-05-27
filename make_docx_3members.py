import sys
import subprocess

try:
    import docx
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    import docx
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = docx.Document()

# Title
title = doc.add_heading('ECOZONE Fingerprint Attendance System', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle = doc.add_heading('OJT Defense Presentation Script — 3 Members', level=1)
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph('Members: Speaker 1 | Speaker 2 | Speaker 3').alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph('─' * 60)

# ----- SLIDE 1 -----
doc.add_heading('SLIDE 1: Introduction & The Problem', level=2)
doc.add_paragraph('[Speaker 1]', style='Intense Quote')
doc.add_paragraph(
    '"Good morning, panelists and faculty. I am [Speaker 1], and together with my teammates [Speaker 2] and [Speaker 3], '
    'we will be presenting our OJT capstone project: the ECOZONE Fingerprint Attendance System.\n\n'
    'During our On-The-Job Training, we observed that the company relied on manual or PIN-based attendance logging, '
    'which is prone to buddy-punching and data inaccuracies. Our solution was to develop a fully hardware-integrated '
    'biometric attendance system that enforces strict fingerprint uniqueness — making it impossible for one employee '
    'to register or clock in on behalf of another."'
)

# ----- SLIDE 2 -----
doc.add_heading('SLIDE 2: OJT Timeline & System Architecture', level=2)
doc.add_paragraph('[Speaker 2]', style='Intense Quote')
doc.add_paragraph(
    '"Thank you, [Speaker 1]. I will walk you through our OJT timeline and how the system is architected.\n\n'
    'Our development was broken into three phases:'
)
doc.add_paragraph("Weeks 1–2: Hardware Familiarization — We studied the DigitalPersona U.are.U fingerprint scanner SDK and set up the C++ development environment.", style='List Bullet')
doc.add_paragraph("Weeks 3–4: API Bridge Development — We built a bridge between our low-level C++ scanner executable and a PHP/MySQL web backend, allowing the web frontend to trigger hardware scans.", style='List Bullet')
doc.add_paragraph("Weeks 5–6: Security Hardening — We implemented real-time duplicate detection both across the database and within a live 10-finger enrollment session.", style='List Bullet')
doc.add_paragraph(
    '\nArchitecturally, the system has three layers: the C++ hardware layer (scanner.exe), the PHP API layer, '
    'and the HTML/JavaScript frontend. These communicate in a pipeline — the frontend calls the API, the API '
    'triggers the scanner, and the result flows back up."'
)

# ----- SLIDE 3 -----
doc.add_heading('SLIDE 3: Technical Deep Dive — Biometric Feature Extraction', level=2)
doc.add_paragraph('[Speaker 3]', style='Intense Quote')
doc.add_paragraph(
    '"Thank you, [Speaker 2]. I will now explain the most critical technical component: how we process and store the fingerprint.\n\n'
    'Panelists often ask: what kind of hashing do we use for the fingerprints? The honest and precise answer is: '
    'we do not use traditional cryptographic hashing such as MD5 or SHA-256.\n\n'
    'Instead, our C++ program uses the ANSI 378 Minutiae Feature Extraction algorithm provided by the DigitalPersona SDK. '
    'Think of it as a cartographer drawing a map. The algorithm identifies the unique ridge endings, bifurcations, and '
    'whorls of the finger — called minutiae — and encodes their X/Y coordinates and angles into a compact binary template.\n\n'
    'This binary template is then converted into a Hexadecimal string by our BytesToHex() function in scanner.cpp, '
    'and stored in the MySQL database as plain text. During login verification, the scanned finger is compared against '
    'all stored templates using a dissimilarity score algorithm — not an exact string match — which tolerates minor '
    'placement differences while still rejecting imposters."'
)

# ----- SLIDE 4 -----
doc.add_heading('SLIDE 4: Key Challenge — Real-Time Session Duplicate Detection', level=2)
doc.add_paragraph('[Speaker 1]', style='Intense Quote')
doc.add_paragraph(
    '"Our biggest engineering challenge was not storing the fingerprints — it was catching duplicates in real time '
    'during a live 10-finger enrollment session.\n\n'
    'The problem: a user scans 10 fingers one by one. If they accidentally place the same finger twice for two '
    'different slots, that data is not yet saved to the database — so a simple database check would miss it.\n\n'
    'Our solution: as each finger is scanned, the PHP backend writes all previously scanned fingers from the current session '
    'into a temporary file. Our C++ executable\'s check mode then compares the new scan against that temp file, '
    'catching duplicates immediately with the message: \'You already scanned this exact fingerprint in the current session!\' '
    'The temp file is securely deleted after each check."'
)

# ----- SLIDE 5 -----
doc.add_heading('SLIDE 5: Live System Demonstration', level=2)
doc.add_paragraph('[Speaker 2 — controls demo; Speaker 3 narrates]', style='Intense Quote')
doc.add_paragraph('Speaker 3: "We will now walk through a live demonstration of the system."')
doc.add_paragraph("[Speaker 2 opens the enrollment page] Speaker 3: \"Here is our enrollment interface. An employee enters their nickname to begin.\"", style='List Number')
doc.add_paragraph("[Speaker 2 places finger on scanner] Speaker 3: \"The system immediately communicates with the hardware SDK and captures the fingerprint. The UI updates in real time.\"", style='List Number')
doc.add_paragraph("[Speaker 2 places the SAME finger in the next slot] Speaker 3: \"Notice the Security Alert — our session validator caught that this fingerprint was already scanned!\"", style='List Number')
doc.add_paragraph("[Speaker 2 completes all 10 fingers and saves] Speaker 3: \"Upon saving, all 10 Minutiae templates are committed to the database. The enrollment is complete.\"", style='List Number')
doc.add_paragraph("[Speaker 2 opens attendance/verify page and scans] Speaker 3: \"And here, the system matches a live scan against all stored templates and identifies the employee instantly.\"", style='List Number')

# ----- SLIDE 6 -----
doc.add_heading('SLIDE 6: Conclusion & OJT Learnings', level=2)
doc.add_paragraph('[Speaker 1]', style='Intense Quote')
doc.add_paragraph(
    '"This OJT project gave our team hands-on experience that goes far beyond what textbooks can provide. '
    '[Speaker 1] gained expertise in C++ systems programming and hardware SDK integration. '
    '[Speaker 2] grew in PHP backend development and API design. '
    '[Speaker 3] developed deep understanding of biometric security standards and database architecture.\n\n'
    'The final deliverable is a fully functioning, secure, duplicate-proof ECOZONE Fingerprint Attendance System '
    'ready for real-world deployment. We are proud of what we accomplished and are ready to answer your questions. '
    'Thank you."'
)

# ----- Q&A GUIDE -----
doc.add_paragraph('─' * 60)
doc.add_heading('Q&A Preparation Guide', level=2)

qa_pairs = [
    (
        "Q: Why didn't you hash the fingerprint data with bcrypt or SHA-256?",
        "A: Cryptographic hashing is designed for fixed text like passwords. For biometrics, even a tiny finger shift creates "
        "a completely different raw image, which would make an exact hash match impossible. Instead, we use the ANSI 378 "
        "Minutiae algorithm, which is a standardized biometric encoding specifically designed to be tolerant of minor variations "
        "while still being secure."
    ),
    (
        "Q: Is the fingerprint data safe if someone accesses the database?",
        "A: Yes. The stored hex strings are binary-encoded Minutiae templates. Without the proprietary DigitalPersona SDK, "
        "they cannot be decoded or reverse-engineered back into a fingerprint image."
    ),
    (
        "Q: What happens if an employee loses a finger?",
        "A: Each employee stores up to 10 fingers. Any one of the remaining fingers can be used to verify attendance. "
        "If a finger slot is unavailable, we allow 'SKIP' during enrollment, requiring at least 1 valid finger."
    ),
    (
        "Q: How accurate is the fingerprint matching?",
        "A: We use a dissimilarity threshold of 21,474, which corresponds to a false positive rate of 0.001% — meaning "
        "only 1 in 100,000 attempts from a different person could be falsely matched."
    ),
]

for q, a in qa_pairs:
    p = doc.add_paragraph()
    run_q = p.add_run(q + "\n")
    run_q.bold = True
    run_a = p.add_run(a)
    run_a.italic = True

output_path = r'C:\Users\Administrator\Desktop\ECOZONE_Defense_Script_3Members.docx'
doc.save(output_path)
print(f"SUCCESS: DOCX created at {output_path}")
