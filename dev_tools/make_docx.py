import sys
import subprocess

try:
    import docx
except ImportError:
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    import docx

doc = docx.Document()
doc.add_heading('ECOZONE Fingerprint Attendance System - Defense Guide', 0)

doc.add_heading('Slide 1: Introduction & The Problem', level=2)
doc.add_paragraph(
    'Speaker 1: "Good morning, panelists. We are here to present our On-The-Job Training (OJT) project: '
    'The ECOZONE Fingerprint Attendance System. During our OJT, we observed that traditional attendance systems '
    'lacked strict physical validation, leading to potential buddy-punching or data inaccuracies. Our goal was to '
    'build a robust, hardware-integrated system that enforces strict biometric uniqueness."'
)

doc.add_heading('Slide 2: The OJT Timeline & Workflow', level=2)
doc.add_paragraph('Speaker 2: "Our OJT journey was structured into clear phases:"')
doc.add_paragraph("Weeks 1-2: Familiarization with the DigitalPersona SDK and C++ hardware drivers.", style='List Bullet')
doc.add_paragraph("Weeks 3-4: Building the bridging API. We successfully linked a C++ executable with a PHP/Web frontend.", style='List Bullet')
doc.add_paragraph("Weeks 5-6: Focusing on security edge cases — preventing duplicate records within a live scanning session and optimizing 10-finger profile insertion.", style='List Bullet')

doc.add_heading('Slide 3: Deep Technical Dive (30% of Score)', level=2)
doc.add_paragraph(
    'Speaker 1: "One of our most complex decisions was how to handle the fingerprint data securely. '
    'We do not use traditional cryptographic hashes like MD5 or bcrypt. Instead, our C++ scanner uses the '
    'ANSI 378 Minutiae Feature Extraction algorithm — which maps the physical ridges and loops of the finger. '
    'The C++ executable serializes this map into a raw Hexadecimal string, which the backend stores directly. '
    'During verification, the system computes an algorithmic dissimilarity score — not an exact hash match — '
    'to ensure accuracy even if a user places their finger at a slight angle."'
)

doc.add_heading('Slide 4: Challenges & Solutions', level=2)
doc.add_paragraph(
    'Speaker 2: "Our biggest hurdle was real-time session validation. When capturing 10 fingers sequentially, '
    'we had to ensure the same finger was not scanned twice in different slots. Since data is not saved to '
    'the main database until all 10 are done, we implemented a secure temp file system. The C++ module reads '
    'the fingers already in the active session and intercepts duplicates instantly before touching the database."'
)

doc.add_heading('Slide 5: Live Demonstration', level=2)
doc.add_paragraph('Speaker 1: "We will now demonstrate this live."')
doc.add_paragraph("[Action: Open Enrollment Page] Our system allows enrolling an employee with a Nickname.", style='List Number')
doc.add_paragraph("[Action: Scan a finger] Notice the immediate UI feedback when the SDK captures input.", style='List Number')
doc.add_paragraph("[Action: Scan the SAME finger again] Our session-validation triggers a Security Alert: 'You already scanned this exact fingerprint in the current session.'", style='List Number')
doc.add_paragraph("[Action: Finish all 10 fingers and open database] All 10 Minutiae maps are committed to the backend.", style='List Number')

doc.add_heading('Slide 6: Conclusion & OJT Outcomes', level=2)
doc.add_paragraph(
    'Speaker 2: "Through this OJT, our team gained practical experience bridging physical hardware with modern web APIs. '
    'We learned low-level memory handling in C++, stateless session validation in PHP, and biometric security standards. '
    'The result is a fully functioning, duplicate-proof ECOZONE biometric system ready for deployment. '
    'We would be happy to take any questions you have."'
)

doc.add_heading('Tips for Delivery & Q&A', level=2)
doc.add_paragraph("Do not read directly off the slides. Use the script as talking points. Maintain eye contact.", style='List Bullet')
doc.add_paragraph(
    'If asked: "Is your system secure without hashing?" — Answer: "Yes. A traditional hash is useless for biometrics '
    'because any slight change in finger placement changes the image completely. Our ANSI 378 template is heavily '
    'encoded binary data that cannot be reverse-engineered back into a fingerprint without proprietary SDK tools."',
    style='List Bullet'
)

output_path = r'C:\Users\Administrator\Desktop\ECOZONE_Defense_Guide.docx'
doc.save(output_path)
print(f"SUCCESS: DOCX created at {output_path}")
