import sys
import subprocess

try:
    import docx
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    import docx
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = docx.Document()

# --- Title Page ---
title = doc.add_heading('ECOZONE Fingerprint Attendance System', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_heading('OJT Defense Presentation Script', level=1)
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph('Host Company: ZAMBOECOZONE – MIS Division').alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph('Duration: 180 Hours | 11 Weeks  |  BSCS – 2A  |  WMSU').alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph('─' * 60)

team = doc.add_paragraph()
team.alignment = WD_ALIGN_PARAGRAPH.CENTER
team.add_run('Team Members:\n').bold = True
team.add_run('Miko R. Zamora  |  Razel T. Herodias  |  Kevin Trazie C. Librero\n')
team.add_run('IT Officer Trainees — Western Mindanao State University')

doc.add_paragraph('─' * 60)
doc.add_paragraph('NOTE: Replace [Miko], [Razel], [Kevin] with your actual speaking assignments.')
doc.add_page_break()

# ─────────────────────────────────────────────
#  SLIDE 1: Introduction
# ─────────────────────────────────────────────
doc.add_heading('SLIDE 1: Introduction & The Problem', level=2)
doc.add_paragraph('[Miko — Opens the presentation]', style='Intense Quote')
doc.add_paragraph(
    '"Good morning, panelists and faculty. I am Miko Zamora, and together with my teammates Razel Herodias '
    'and Kevin Librero, we are proud to present the culmination of our 180-hour On-The-Job Training at '
    'ZAMBOECOZONE — the Ecozone Fingerprint Attendance System.\n\n'
    'ZAMBOECOZONE, or the Zamboanga City Special Economic Zone Authority and Freeport, is a government-owned '
    'corporate entity and the premier economic hub for trade, manufacturing, and ecotourism in the ASEAN region. '
    'We were assigned as IT Officer Trainees under their Management Information Systems, or MIS, Division.\n\n'
    'During our immersion, we noticed that existing attendance logging relied on manual or PIN-based methods '
    'which are prone to buddy-punching and data inaccuracies. Our goal was to design, develop, and deploy a '
    'hardware-integrated biometric system that enforces strict fingerprint uniqueness — making it impossible '
    'for one employee to record attendance on behalf of another."'
)

# ─────────────────────────────────────────────
#  SLIDE 2: OJT Timeline
# ─────────────────────────────────────────────
doc.add_heading('SLIDE 2: OJT Timeline & System Architecture', level=2)
doc.add_paragraph('[Razel — Presents the weekly breakdown]', style='Intense Quote')
doc.add_paragraph(
    '"Thank you, Miko. I will walk the panel through our 11-week OJT journey and explain how our system was built layer by layer.\n"'
)

timeline = [
    ("Week 1 — Orientation & Onboarding",
     "We joined the official orientation program, toured the HTE building, met our supervisors, and familiarized ourselves with ZAMBOECOZONE's policies and procedures."),
    ("Week 2 — Data Privacy Act Study",
     "We studied the Data Privacy Act of 2012, Republic Act 10173, to ensure our biometric system would comply with all legal obligations for personal data handling and protection."),
    ("Week 3 — Biometric System Kickoff & WordPress CMS",
     "We initiated development of the fingerprint scanner system using C++ and the U.are.U 4500 SDK from DigitalPersona. We implemented biometric feature extraction for secure data storage. We also set up WordPress CMS and explored the Government Web Template."),
    ("Week 4 — CodeIgniter 4 & GitHub Version Control",
     "We set up CodeIgniter 4, learned the MVC architecture, and created a shared GitHub repository with three forks for collaborative development. We made significant backend progress on the biometric project."),
    ("Week 5 — Mid-Point Evaluation (40–50% Complete)",
     "We presented the system to our mentor for evaluation at roughly half-completion. We incorporated feedback and defined the remaining deliverables."),
    ("Week 6 — Deep Dive into CodeIgniter 4",
     "We focused on CI4 fundamentals — routes, controllers, and views — completing small practice projects to solidify the MVC workflow."),
    ("Week 7 — WordPress CMS Exploration",
     "We installed and configured a local WordPress site, created demo pages, customized themes, and resolved plugin-theme conflicts through independent research."),
    ("Week 8 — Network Cable Assembly",
     "We assembled a straight-through Ethernet cable using T568B color coding, connected RJ45 connectors, and verified network connectivity — a hands-on IT support task assigned by MIS."),
    ("Week 9 — Attendance System UI Refinement",
     "We improved the system's UI layouts for better usability, organized project files on GitHub, and clarified the full data flow: from hardware scan input, through API processing, to database storage and dashboard display."),
    ("Week 10 — System Polishing & Final Testing",
     "We polished the system end-to-end — fixing bugs, improving the interface, and conducting full feature testing to prepare for deployment."),
    ("Week 11 — MVC Framework Architecture Study & Wrap-Up",
     "We studied the full CI4 project connection between routes, views, and models. We attended mentoring sessions, completed final team documentation, and submitted our OJT deliverables."),
]

for title_text, description in timeline:
    p = doc.add_paragraph(style='List Bullet')
    run_title = p.add_run(title_text + ': ')
    run_title.bold = True
    p.add_run(description)

doc.add_paragraph(
    '\n[Razel continues]\n"Architecturally, the system operates on three layers: '
    'the C++ hardware layer — scanner.exe — communicates with the U.are.U 4500 fingerprint reader. '
    'The PHP API layer bridges the hardware output to our MySQL database. '
    'And the HTML/JavaScript frontend provides the interface for both enrollment and daily attendance verification. '
    'All three layers work together in a secure pipeline."'
)

# ─────────────────────────────────────────────
#  SLIDE 3: Technical Deep Dive
# ─────────────────────────────────────────────
doc.add_heading('SLIDE 3: Technical Deep Dive — Biometric Feature Extraction', level=2)
doc.add_paragraph('[Kevin — Explains the technical core]', style='Intense Quote')
doc.add_paragraph(
    '"Thank you, Razel. I will now explain the most critical technical component — biometric feature extraction — '
    'which is what our portfolio describes as cryptographic hashing for biometric data.\n\n'
    'When panelists ask what hashing algorithm we used, here is our precise answer: we do not use traditional '
    'cryptographic hashing such as MD5 or SHA-256. For biometrics, those are unsuitable — because even a slight '
    'shift in finger placement changes the raw pixel image entirely, making an exact hash match impossible.\n\n'
    'Instead, our C++ program uses the ANSI 378 Minutiae Feature Extraction algorithm provided by the DigitalPersona '
    'U.are.U 4500 SDK. Think of it as a cartographer drawing a precise map. The algorithm identifies the unique '
    'ridge endings, bifurcations, and whorls on the finger — called minutiae points — and encodes their coordinates '
    'and angles into a compact binary template.\n\n'
    'This binary template is then serialized into a Hexadecimal string by our BytesToHex() function in scanner.cpp, '
    'and stored directly in MySQL. During login verification, the live scan is compared to all stored templates using '
    'a dissimilarity score algorithm — not an exact string match — which tolerates minor angle differences while '
    'still rejecting any other person\'s finger.\n\n'
    'This complies with our RA 10173 training — the biometric data cannot be reverse-engineered into a readable '
    'fingerprint image without the proprietary SDK, making it both legally secure and technically robust."'
)

# ─────────────────────────────────────────────
#  SLIDE 4: Key Challenge
# ─────────────────────────────────────────────
doc.add_heading('SLIDE 4: Key Challenge — Real-Time Session Duplicate Detection', level=2)
doc.add_paragraph('[Miko — Returns to present the challenge]', style='Intense Quote')
doc.add_paragraph(
    '"Our biggest engineering challenge was not simply storing fingerprints — it was catching duplicates in real time '
    'during a live 10-finger enrollment session.\n\n'
    'The problem: a user scans all 10 fingers one by one. If they accidentally place the same finger twice for '
    'two different finger slots, that session data is not yet committed to the database — so a standard database '
    'check would completely miss it.\n\n'
    'Our solution: as each finger is scanned, the PHP backend writes all previously collected fingers from the '
    'current session into a secure temporary file. Our C++ executable\'s check mode then compares the new scan '
    'against that temp file using the same biometric dissimilarity algorithm — intercepting the duplicate before '
    'it reaches the database and displaying the alert: '
    '"You already scanned this exact fingerprint in the current session!" '
    'The temporary file is securely deleted after each check, leaving no trace."'
)

# ─────────────────────────────────────────────
#  SLIDE 5: Live Demo
# ─────────────────────────────────────────────
doc.add_heading('SLIDE 5: Live System Demonstration', level=2)
doc.add_paragraph('[Kevin — Controls the demo  |  Razel — Narrates]', style='Intense Quote')
doc.add_paragraph('Razel: "We will now walk through a live demonstration of the system in action."')
doc.add_paragraph(
    '[Kevin opens the enrollment page]  Razel: "Here is our enrollment interface. The employee enters their full name to begin the 10-finger capture process."',
    style='List Number'
)
doc.add_paragraph(
    '[Kevin places a finger on the scanner]  Razel: "The system calls our PHP API, which triggers scanner.exe in the background. Notice the real-time UI feedback as the fingerprint is captured."',
    style='List Number'
)
doc.add_paragraph(
    '[Kevin places the SAME finger in the next slot]  Razel: "Watch carefully — our session validator immediately triggers a Security Alert: \'You already scanned this exact fingerprint.\' No duplicate makes it into the database."',
    style='List Number'
)
doc.add_paragraph(
    '[Kevin completes all 10 fingers and clicks Save]  Razel: "Upon final save, all 10 ANSI 378 Minutiae templates are committed to the database in a single secure transaction."',
    style='List Number'
)
doc.add_paragraph(
    '[Kevin opens the attendance/verify page and scans]  Razel: "During daily attendance, the live scan is matched against all stored profiles. The system identifies the employee and logs the timestamp instantly."',
    style='List Number'
)

# ─────────────────────────────────────────────
#  SLIDE 6: Conclusion
# ─────────────────────────────────────────────
doc.add_heading('SLIDE 6: Conclusion, Reflections & OJT Learnings', level=2)
doc.add_paragraph('[All three — Razel leads, Miko and Kevin add their personal reflection]', style='Intense Quote')
doc.add_paragraph(
    'Razel: "This OJT experience at ZAMBOECOZONE was transformative for all three of us. Building this system '
    'from scratch — integrating hardware SDK, backend logic, and a web interface — gave us a complete view of '
    'how real-world IT systems are built within professional organizations.\n\n'
    'Miko: "For me personally, the biggest growth was in C++ systems programming and hardware integration. '
    'Learning how to bridge a low-level scanner DLL with a web-accessible API was a challenge I had never '
    'faced in the classroom.\n\n'
    'Kevin: "I developed a deep understanding of biometric security standards, the Data Privacy Act, and how '
    'to design a system that is both user-friendly and legally compliant.\n\n'
    'Razel: "Over 11 weeks, we grew from orientation to deploying a fully functioning, duplicate-proof biometric '
    'attendance system. Working within the MIS Division also sharpened our professional skills — communication, '
    'collaboration, and receiving constructive feedback gracefully.\n\n'
    'The final system is ready for real-world deployment. We are proud of what we built as Team ECOZONE, '
    'and we are ready to take any questions the panel may have. Thank you."'
)

# ─────────────────────────────────────────────
#  Q&A Guide
# ─────────────────────────────────────────────
doc.add_page_break()
doc.add_heading('Q&A Preparation Guide', level=2)
doc.add_paragraph('Use these model answers during the panel Q&A:')

qa_pairs = [
    (
        "Q: What hashing algorithm did you use for the fingerprints?",
        "A: We use the ANSI 378 Minutiae Feature Extraction algorithm via the DigitalPersona U.are.U 4500 SDK — not traditional cryptographic hashing. It maps the physical geometry of the finger's ridges and encodes them as a binary template, which our BytesToHex() function converts to a storable string. During verification, a dissimilarity score is computed rather than an exact match."
    ),
    (
        "Q: Is the stored fingerprint data safe under RA 10173?",
        "A: Yes. The stored hex strings are encoded ANSI 378 binary templates. They cannot be converted back into a fingerprint image without the proprietary SDK. No raw biometric image is ever stored, which aligns with RA 10173 data minimization and security principles."
    ),
    (
        "Q: How accurate is the matching? What is the false positive rate?",
        "A: Our system uses a dissimilarity threshold of 21,474, which corresponds to a false positive rate of 0.001% — or 1 in 100,000 verification attempts from a different person could be falsely matched."
    ),
    (
        "Q: What happens if an employee is missing a finger or their scan fails?",
        "A: Each employee registers up to 10 fingers. Any enrolled finger can be used for verification. If a slot cannot be scanned — due to injury or hardware error — our system allows the slot to be marked as 'SKIP', with a minimum of at least 1 valid fingerprint required for registration."
    ),
    (
        "Q: What technologies did you use overall?",
        "A: C++ with the DigitalPersona U.are.U 4500 SDK for hardware integration, PHP and MySQL for the backend API and database, HTML/CSS/JavaScript for the web frontend, CodeIgniter 4 for MVC framework study, and GitHub for collaborative version control."
    ),
]

for q, a in qa_pairs:
    p = doc.add_paragraph()
    run_q = p.add_run(q + '\n')
    run_q.bold = True
    run_a = p.add_run(a)
    run_a.italic = True

output_path = r'C:\Users\Administrator\Desktop\ECOZONE_Defense_Script_Final.docx'
doc.save(output_path)
print(f"SUCCESS: DOCX saved to {output_path}")
