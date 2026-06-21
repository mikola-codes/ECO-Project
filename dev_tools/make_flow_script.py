import sys
import subprocess

try:
    import docx
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    import docx
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = docx.Document()

def heading(text, level=2):
    return doc.add_heading(text, level=level)

def speaker(name, color_marker=""):
    p = doc.add_paragraph()
    r = p.add_run(f"[{name}]")
    r.bold = True
    return p

def body(text):
    return doc.add_paragraph(text)

def bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        r = p.add_run(bold_prefix + ": ")
        r.bold = True
    p.add_run(text)
    return p

def divider():
    doc.add_paragraph("─" * 65)

# ══════════════════════════════════════════
#  TITLE PAGE
# ══════════════════════════════════════════
t = doc.add_heading("ECOZONE Fingerprint Attendance System", 0)
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
s = doc.add_heading("System Flow Defense Script  |  3 Members", level=1)
s.alignment = WD_ALIGN_PARAGRAPH.CENTER
info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
info.add_run("ZAMBOECOZONE – MIS Division  |  180 Hours  |  BSCS-2A  |  WMSU\n").bold = True
info.add_run("Miko R. Zamora   |   Razel T. Herodias   |   Kevin Trazie C. Librero")
divider()
note = doc.add_paragraph()
note.add_run("HOW TO USE THIS SCRIPT: ").bold = True
note.add_run(
    "Each section shows which member speaks. Speak naturally — do NOT read word for word. "
    "Use this as your detailed talking guide. Italicized lines are ACTION cues during the live demo."
)
doc.add_page_break()

# ══════════════════════════════════════════
#  PART 1: MIKO — System Overview & Architecture
# ══════════════════════════════════════════
heading("PART 1: System Overview & Architecture", level=1)
speaker("MIKO — Opens, presents system overview and architecture")
divider()

body(
    '"Good morning, panelists. I am Miko Zamora. I will start by giving you the complete picture of what '
    'our system is, what problem it solves, and how it is built end-to-end.\n'
)

heading("What is the System?")
body(
    'The ECOZONE Fingerprint Attendance System is a hardware-integrated biometric attendance tracker built '
    'for ZAMBOECOZONE\'s MIS Division. It replaces manual or PIN-based attendance logging with physical '
    'fingerprint verification — ensuring that only the registered employee can record their own attendance.'
)

heading("The Problem It Solves")
bullet("Buddy-punching — one employee clocking in for another", "Problem 1")
bullet("Data inaccuracy from manual logs or shared PIN codes", "Problem 2")
bullet("No real-time duplicate protection during multi-finger enrollment", "Problem 3")

heading("System Architecture — 3 Layers")
body("Our system is a three-layer pipeline that connects physical hardware to a web interface:")

bullet(
    "scanner.exe — A compiled C++ executable that directly communicates with the "
    "DigitalPersona U.are.U 4500 fingerprint scanner hardware via the dpfpdd.dll and dpfj.dll SDK libraries. "
    "It has three modes: ENROLL (capture a finger and return data), VERIFY (scan and match against a file), "
    "and CHECK (compare a given hex string against a file without scanning).",
    "Layer 1 — C++ Hardware Layer"
)
bullet(
    "PHP API (XAMPP/Apache) — Receives requests from the frontend. It calls scanner.exe using PHP's "
    "shell_exec() function, handles the fingerprint hex data, performs duplicate checks, and executes all "
    "MySQL database reads and writes using prepared statements.",
    "Layer 2 — PHP Backend API"
)
bullet(
    "HTML/CSS/JavaScript Frontend — The web interface the user sees. It presents the enrollment form, "
    "calls the API for each finger scan, collects all 10 responses, and then sends the final save request. "
    "It also calls the attendance/verify API for daily clock-ins.",
    "Layer 3 — Web Frontend"
)

body(
    '"These three layers communicate in a strict pipeline. The web page cannot talk directly to the scanner — '
    'it always goes through the PHP API, which in turn calls the C++ executable. This separation of concerns '
    'makes the system secure, maintainable, and testable at each layer independently."'
)

doc.add_page_break()

# ══════════════════════════════════════════
#  PART 2: RAZEL — Enrollment Flow
# ══════════════════════════════════════════
heading("PART 2: Enrollment Flow — Step by Step", level=1)
speaker("RAZEL — Explains the full enrollment process, with live demo narration")
divider()

body(
    '"Thank you, Miko. I am Razel Herodias. I will now walk through the complete fingerprint enrollment flow '
    '— from the moment a user opens the enrollment page to the moment their 10 fingerprints are saved in the database.\n'
)

heading("Step 1 — User Opens Enrollment Page")
body(
    'The user opens index.html and fills in their Nickname. The frontend validates that the name is not empty '
    'and does not exceed 100 characters. Before any scanning begins, the PHP API checks the database for a '
    'duplicate nickname using a prepared statement — if the name is taken, enrollment is blocked immediately.'
)

heading("Step 2 — 10-Finger Sequential Scanning Loop")
body(
    'The frontend runs a loop for fingers 0 through 9 — Right Thumb, Right Index, Right Middle, Right Ring, '
    'Right Pinky, Left Thumb, Left Index, Left Middle, Left Ring, and Left Pinky — in that exact order.'
)
body('For each finger, the frontend sends an API call to enroll.php with the finger_index parameter. The PHP server then:')
bullet('Calls scanner.exe enroll via shell_exec(), which activates the hardware scanner.')
bullet('scanner.exe opens the USB device, captures the raw fingerprint image into a 500,000-byte buffer.')
bullet('It extracts the ANSI 378 Minutiae Feature Template using the dpfj_create_fmd_from_raw() SDK function.')
bullet('It converts the binary template to a Hexadecimal string using the BytesToHex() function.')
bullet('This hex string is returned through stdout back to PHP, and PHP echoes it as a JSON response to the frontend.')

heading("Step 3 — Real-Time Session Duplicate Check (Critical Security Feature)")
body(
    'This is the most important security step. Before accepting each finger\'s hex data, the PHP server performs '
    'a cross-check against two sources:'
)
bullet(
    'All previously registered employees — PHP queries the full fingerprints table and writes all their stored '
    'hex data into a temporary .tmp file on the server.',
    "Source A"
)
bullet(
    'All fingers already scanned in the current session — The frontend sends back every previously collected '
    'finger as session_finger_0 through session_finger_9 POST parameters. PHP adds these into the same temp file.',
    "Source B"
)
body(
    'PHP then calls scanner.exe in CHECK mode, passing the newly scanned hex string and the temp file path. '
    'scanner.exe runs the dpfj_compare() SDK function — computing a biometric dissimilarity score — against every '
    'record in the file. If any score falls below the threshold of 21,474 (false positive rate: 0.001%), '
    'it returns MATCH:[id]. PHP immediately rejects the scan and returns a Security Alert to the user. '
    'The temp file is deleted after every single check.'
)

heading("Step 4 — User Can Skip a Finger")
body(
    'If an employee is missing a finger or cannot scan successfully, the user clicks Skip. '
    'The frontend sends fingerprint_data=SKIP for that slot. At least 1 real fingerprint is required — '
    'the system blocks saving if all 10 are skipped.'
)

heading("Step 5 — Final Save to Database")
body(
    'After all 10 fingers are collected (as hex strings or SKIP), the frontend sends one final POST request '
    'to enroll.php containing all 10 values as finger_0 through finger_9. PHP uses a single prepared INSERT '
    'statement to save the record into the fingerprints table with 11 columns: nickname and one column per finger.'
)

body(
    '"The result is a complete, duplicate-proof employee biometric profile stored securely in MySQL. '
    'Each hex string in the database is an encoded ANSI 378 template — unreadable without the SDK."'
)

# LIVE DEMO CUE
p = doc.add_paragraph()
p.add_run("\n[LIVE DEMO — RAZEL narrates, MIKO controls the screen]\n").bold = True
doc.add_paragraph('1. Open the enrollment page. Type a Nickname and click Enroll.', style='List Number')
doc.add_paragraph('2. Place a finger. Show the real-time UI updating as the scan completes.', style='List Number')
doc.add_paragraph('3. Place the SAME finger in the next slot. Show the Security Alert popup.', style='List Number')
doc.add_paragraph('4. Complete all 10 fingers. Click Save. Show the success message.', style='List Number')
doc.add_paragraph('5. Open phpMyAdmin or show the database — show the hex strings stored per finger.', style='List Number')

doc.add_page_break()

# ══════════════════════════════════════════
#  PART 3: KEVIN — Verification Flow & Technical Security
# ══════════════════════════════════════════
heading("PART 3: Verification Flow & Technical Security", level=1)
speaker("KEVIN — Explains the attendance verification process and biometric security design")
divider()

body(
    '"Thank you, Razel. I am Kevin Librero. I will now explain how the system verifies a fingerprint '
    'during daily attendance, and why this system is technically secure.\n'
)

heading("Verification Flow — Daily Attendance Clock-In")
body('When an employee arrives and needs to record attendance, the following steps occur:')

bullet(
    'The attendance page calls verify.php on the backend. PHP calls scanner.exe verify [tempfile] via shell_exec(). '
    'The scanner hardware captures the live fingerprint — one scan, no retry loop.',
    "Step 1 — Live Scan Triggered"
)
bullet(
    'Before calling the scanner, PHP queries the full fingerprints table and writes every employee\'s ID and '
    'all their 10 stored hex strings into a temporary file in the format: employee_id|fmd1|fmd2|...|fmd10.',
    "Step 2 — Database Loaded into Temp File"
)
bullet(
    'scanner.exe reads the temp file line by line. For each employee, it runs dpfj_compare() between the live '
    'scan and each of that employee\'s stored finger templates. It tracks the lowest dissimilarity score seen. '
    'If any score is below 21,474, that employee is flagged as a match.',
    "Step 3 — 1-to-N Biometric Matching"
)
bullet(
    'scanner.exe outputs either MATCH:[employee_id] or NOMATCH to stdout. PHP reads this and returns '
    'the result as JSON to the frontend. If matched, PHP logs the attendance timestamp to the attendance table.',
    "Step 4 — Result Logged"
)

heading("Why 21,474 as the Threshold?")
body(
    'The DigitalPersona SDK documentation defines this threshold as corresponding to a False Accept Rate (FAR) '
    'of 0.001% — meaning only 1 in 100,000 attempts by a different person would be falsely accepted. '
    'A lower score means more similarity. Score = 0 is a perfect match. Scores above 21,474 are rejected.'
)

heading("Why Not MD5 or SHA-256?")
body(
    'Traditional cryptographic hashing is a one-way function for fixed digital data — like passwords. '
    'For fingerprints, it is completely unsuitable: even a 1-degree rotation of the same finger produces '
    'an entirely different byte array, making any exact hash comparison always fail. '
    'The ANSI 378 Minutiae algorithm was specifically designed for biometrics — it encodes the structural '
    'geometry of the finger, not pixel values, making it inherently tolerant of minor placement variations.'
)

heading("Data Privacy Compliance — RA 10173")
bullet('No raw fingerprint image is ever stored. Only the extracted ANSI 378 template is saved.')
bullet('The hex string in the database cannot be reverse-engineered into a fingerprint without the proprietary SDK.')
bullet('Temporary files used during checks are immediately deleted after each operation.')
bullet('All database queries use prepared statements — immune to SQL injection.')

heading("Technologies Used")
bullet("C++ with DigitalPersona U.are.U 4500 SDK (dpfpdd.dll, dpfj.dll) — hardware integration", "Hardware")
bullet("PHP 8 + MySQL via MySQLi prepared statements — backend API", "Backend")
bullet("HTML5 / CSS3 / JavaScript — web frontend interface", "Frontend")
bullet("XAMPP / Apache — local web server environment", "Server")
bullet("GitHub — collaborative version control with 3 forks on a shared branch", "Version Control")
bullet("CodeIgniter 4 — studied for MVC architecture during OJT (parallel task)", "Framework")

body(
    '"In summary: the system captures a physical biometric input, extracts a mathematically unique template, '
    'stores it securely, and matches it in real time using a proven SDK algorithm — all while complying with '
    'the Data Privacy Act and rejecting duplicates at every stage of the pipeline."'
)

# LIVE DEMO CUE
p = doc.add_paragraph()
p.add_run("\n[LIVE DEMO — KEVIN controls screen]\n").bold = True
doc.add_paragraph('1. Open the attendance/verify page.', style='List Number')
doc.add_paragraph('2. Place an enrolled finger on the scanner.', style='List Number')
doc.add_paragraph('3. Show the MATCH result — employee name identified, timestamp logged.', style='List Number')
doc.add_paragraph('4. Place an unregistered finger. Show NOMATCH result.', style='List Number')

doc.add_page_break()

# ══════════════════════════════════════════
#  CLOSING — ALL THREE
# ══════════════════════════════════════════
heading("CLOSING STATEMENT", level=1)
speaker("MIKO — Closes on behalf of all three")
divider()
body(
    '"To summarize the system flow: A user\'s fingerprint is captured in C++ via the scanner SDK. '
    'The raw image is processed into an ANSI 378 Minutiae template. That template is serialized to a hex string '
    'and stored in MySQL. Every new scan is checked against all stored templates and all in-session fingers '
    'using biometric dissimilarity scoring — blocking any duplicate before it touches the database. '
    'During daily attendance, the system performs a 1-to-N match across the entire employee database and '
    'returns a result in seconds.\n\n'
    'This is the ECOZONE Fingerprint Attendance System. We are Miko, Razel, and Kevin — thank you, '
    'and we are ready for your questions."'
)

doc.add_page_break()

# ══════════════════════════════════════════
#  Q&A CHEAT SHEET
# ══════════════════════════════════════════
heading("Q&A CHEAT SHEET", level=1)
body("Memorize these. Answer confidently. Do not guess.")
divider()

qa = [
    ("What hashing did you use?",
     "We use ANSI 378 Minutiae Feature Extraction via the DigitalPersona SDK — not traditional cryptographic hashing. "
     "It maps ridge geometry, not pixels, so it tolerates minor finger placement variation."),
    ("Is the data secure under RA 10173?",
     "Yes. We never store a raw fingerprint image. The ANSI 378 hex template cannot be reversed without the proprietary SDK. "
     "Temp files are deleted after each check. All SQL uses prepared statements."),
    ("What is the FAR / accuracy?",
     "Our threshold of 21,474 gives a False Accept Rate of 0.001% — 1 in 100,000 cross-person attempts would be falsely accepted."),
    ("How does session duplicate detection work?",
     "Each finger scanned is passed back from the frontend as a POST parameter. PHP writes all session fingers + all database "
     "records into a temp file. C++ runs biometric comparison against the entire temp file. If any match is found, it is rejected immediately."),
    ("What if an employee is missing a finger?",
     "They can SKIP that slot. At minimum 1 real fingerprint is required. Any enrolled finger can verify attendance."),
    ("Why did you use C++ instead of doing everything in PHP?",
     "The DigitalPersona SDK is a native Windows DLL (dpfpdd.dll, dpfj.dll). It can only be called from a compiled C++ application. "
     "PHP cannot load native DLLs directly, so we bridge the two via shell_exec()."),
    ("What was your biggest technical challenge?",
     "Real-time session duplicate detection. The data isn't in the database yet during a session, so we had to implement "
     "a dynamic temp file system to check in-progress scans against both the database and the current session simultaneously."),
]

for q, a in qa:
    p = doc.add_paragraph()
    p.add_run("Q: " + q + "\n").bold = True
    r = p.add_run("A: " + a)
    r.italic = True

output_path = r'C:\Users\Administrator\Desktop\ECOZONE_SystemFlow_Script.docx'
doc.save(output_path)
print(f"SUCCESS: DOCX saved to {output_path}")
